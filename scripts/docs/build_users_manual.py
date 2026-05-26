from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = PROJECT_ROOT / "storage" / "output"
OUTPUT_PATH = OUTPUT_DIR / "Users Manual.docx"
ASSET_DIR = PROJECT_ROOT / "storage" / "tmp" / "users-manual-assets" / "annotated"
LOGO_PATH = PROJECT_ROOT / "assets" / "Logo.png"
MANUAL_DATE = "May 25, 2026"


BODY_FONT = "Calibri"
TITLE_COLOR = RGBColor(32, 55, 72)
HEADING_1_COLOR = RGBColor(46, 116, 181)
HEADING_2_COLOR = RGBColor(31, 77, 120)
MUTED_TEXT = RGBColor(85, 85, 85)
TABLE_HEADER_FILL = "E8EEF5"
TABLE_BORDER = "AFC4DA"
NOTE_FILL = "F4F6F9"
NOTE_BORDER = "D5DDE7"

CONTENT_WIDTH = Inches(6.5)
CALL_OUT_COL_1 = Inches(1.181)
CALL_OUT_COL_2 = Inches(5.319)
STEP_COL_1 = Inches(0.9)
STEP_COL_2 = Inches(5.6)


@dataclass(frozen=True)
class ManualSection:
    title: str
    intro: tuple[str, ...]
    image_name: str
    figure_title: str
    callouts: tuple[tuple[str, str], ...]
    steps: tuple[tuple[str, str], ...] = ()


SECTIONS: tuple[ManualSection, ...] = (
    ManualSection(
        title="1. Sign In",
        intro=(
            "Open the DTMIS sign-in page and enter your assigned username and password.",
            "If you cannot sign in, use the password recovery link or contact your system administrator.",
        ),
        image_name="01-login-annotated.png",
        figure_title="Figure 1. Sign-in screen",
        callouts=(
            ("1", "Username field. Type the username assigned to your account."),
            ("2", "Password field. Enter your current password."),
            ("3", "Remember me. Use this only on a secure work device."),
            ("4", "Forgot password. Opens the recovery process when you cannot sign in."),
            ("5", "Sign In button. Starts your session and sends you to your role dashboard."),
        ),
        steps=(
            ("1", "Open the system login page."),
            ("2", "Type your username and password."),
            ("3", "Click Sign In."),
        ),
    ),
    ManualSection(
        title="2. Dashboard Overview",
        intro=(
            "The dashboard is your starting point after sign-in.",
            "It shows your role menu, current workload summary, and quick access to daily actions.",
        ),
        image_name="02-dashboard-annotated.png",
        figure_title="Figure 2. Dashboard overview",
        callouts=(
            ("1", "Create / Intake menu. Opens the page where new documents are encoded."),
            ("2", "Notifications bell. Shows routing updates and reminders."),
            ("3", "Theme switch. Changes the screen theme for comfort and visibility."),
            ("4", "Profile menu. Opens profile settings, screenshot tools, and logout."),
            ("5", "Workflow Gatekeeping panel. Summarizes documents waiting for action."),
            ("6", "ARTA Risk Monitor. Highlights due-soon or overdue workload."),
        ),
    ),
    ManualSection(
        title="3. Notifications Panel",
        intro=(
            "Use notifications to monitor document movement without searching every record manually.",
            "This is useful when you want a quick update on recently routed or received documents.",
        ),
        image_name="03-notifications-annotated.png",
        figure_title="Figure 3. Notifications and alerts",
        callouts=(
            ("1", "Notifications bell. Click to open or close the alert panel."),
            ("2", "Open page. Moves to the full notifications page for longer review."),
            ("3", "Mark all read. Clears unread indicators after you review the items."),
            ("4", "Notification item. Opens the related document or tracking slip when available."),
        ),
    ),
    ManualSection(
        title="4. Create / Intake Workspace",
        intro=(
            "Use the Create / Intake page to review recent intake records and open the intake form.",
            "This page also provides quick document tools, filters, and tracking-slip access.",
        ),
        image_name="04-create-intake-page-annotated.png",
        figure_title="Figure 4. Create / Intake page",
        callouts=(
            ("1", "Create / Intake menu. Confirms that you are in the encoding workspace."),
            ("2", "Create Intake button. Opens the document intake form."),
            ("3", "Status filter. Narrows the table to the queue status you want to review."),
            ("4", "Recent Intakes section. Lists the latest encoded documents."),
            ("5", "View Tracking Slip. Opens the slip for the selected record."),
        ),
        steps=(
            ("1", "Click Create / Intake in the left menu."),
            ("2", "Click Create Intake."),
            ("3", "Fill in the form, then submit the document."),
            ("4", "Return to Recent Intakes and open the tracking slip to confirm the record."),
        ),
    ),
    ManualSection(
        title="5. Create Document Intake Form",
        intro=(
            "The intake form is where you encode a new document before it is tracked inside the system.",
            "Complete the required fields carefully so the receiving office sees the correct details.",
        ),
        image_name="05-create-intake-modal-annotated.png",
        figure_title="Figure 5. Create Document Intake form",
        callouts=(
            ("1", "Sender. Enter the external sender name when applicable."),
            ("2", "Subject / Title. Type the main document title or subject."),
            ("3", "Document Type. Choose the correct document classification."),
            ("4", "Attachments. Upload the digital copy when available."),
            ("5", "Remarks. Add extra notes for the log or receiving office."),
            ("6", "Submit Intake. Saves the intake and generates the tracking-ready record."),
        ),
    ),
    ManualSection(
        title="6. Tracking Slip",
        intro=(
            "The tracking slip is the main proof of document movement inside DTMIS.",
            "Use it to review the tracking ID, origin, subject, and the custody timeline.",
        ),
        image_name="06-tracking-slip-annotated.png",
        figure_title="Figure 6. Tracking slip view",
        callouts=(
            ("1", "Back button. Returns you to the previous page."),
            ("2", "Tracking ID field. Enter another tracking number when you want to load a different record."),
            ("3", "Load Tracking Slip. Opens the slip for the tracking ID in the field."),
            ("4", "Print Slip. Prints the current tracking slip."),
            ("5", "Print Full Package. Opens the larger print package with more record details."),
            ("6", "Tracking slip body. Shows the official slip information and routing timeline."),
        ),
        steps=(
            ("1", "Open a record from Recent Intakes or another queue page."),
            ("2", "Review the tracking ID, subject, and offices shown in the slip."),
            ("3", "Use Print Slip or Print Full Package when a hard copy is needed."),
        ),
    ),
    ManualSection(
        title="7. Search Completed Records",
        intro=(
            "Use the search/archive page to find completed records under your office scope.",
            "This is helpful for historical review, retrieval, and audit preparation.",
        ),
        image_name="07-search-documents-annotated.png",
        figure_title="Figure 7. Search completed records",
        callouts=(
            ("1", "Archive area. Main workspace for reviewing completed documents."),
            ("2", "Date filter. Narrows the archive to a selected date range."),
            ("3", "Completed Document Archive. The results table for completed records."),
            ("4", "View Details. Opens the selected document details panel."),
        ),
    ),
    ManualSection(
        title="8. Profile Settings",
        intro=(
            "Use profile settings to confirm your account details and maintain your access.",
            "You can update your email address and change your password from this page.",
        ),
        image_name="08-profile-settings-annotated.png",
        figure_title="Figure 8. Profile settings page",
        callouts=(
            ("1", "Account summary. Shows your name, username, and current role."),
            ("2", "Update Email section. Starts the email update process."),
            ("3", "Email field. Enter the replacement Gmail or DENR email address."),
            ("4", "Change Password section. Opens the password update form."),
            ("5", "New password field. Type the password you want to use next."),
        ),
        steps=(
            ("1", "Open Profile Settings from the sidebar or profile menu."),
            ("2", "Update your email or password, depending on your need."),
            ("3", "Enter your current password when the page asks for verification."),
        ),
    ),
)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_margins = tc_pr.first_child_found_in("w:tcMar")
    if tc_margins is None:
        tc_margins = OxmlElement("w:tcMar")
        tc_pr.append(tc_margins)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_border(paragraph, color: str, size: int = 12) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color: str = TABLE_BORDER, size: str = "8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_inches: float = 6.5, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_column_widths(table, widths: Iterable[float]) -> None:
    rows = table.rows
    widths = list(widths)
    for row in rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def apply_run_font(run, name=BODY_FONT, size=11, color: RGBColor | None = None, bold=False, italic=False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, HEADING_1_COLOR, 18, 10),
        ("Heading 2", 13, HEADING_1_COLOR, 14, 7),
        ("Heading 3", 12, HEADING_2_COLOR, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_para.paragraph_format.space_after = Pt(3)
    header_run = header_para.add_run("DENR Region XII DTMIS | Users Manual")
    apply_run_font(header_run, size=9, color=MUTED_TEXT, bold=True)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.paragraph_format.space_before = Pt(3)
    footer_run = footer_para.add_run("Operational guide for role-based system users")
    apply_run_font(footer_run, size=8, color=MUTED_TEXT)


def add_cover(doc: Document) -> None:
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(22)
    if LOGO_PATH.exists():
        cover.add_run().add_picture(str(LOGO_PATH), width=Inches(1.15))

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(18)
    kicker.paragraph_format.space_after = Pt(6)
    kicker_run = kicker.add_run("DENR REGION XII DTMIS")
    apply_run_font(kicker_run, size=11, color=RGBColor(122, 90, 0), bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Users Manual")
    apply_run_font(title_run, size=26, color=TITLE_COLOR, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("How to use the role-based Document Tracking and Monitoring Information System")
    apply_run_font(subtitle_run, size=13, color=RGBColor(43, 81, 99))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(0)
    meta_run = meta.add_run(f"Sample walkthrough role: CENRO Admin Record | Updated {MANUAL_DATE}")
    apply_run_font(meta_run, size=10.5, color=MUTED_TEXT, bold=True)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(16)
    spacer.paragraph_format.space_after = Pt(8)
    set_paragraph_border(spacer, "D5DDE7")

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.space_before = Pt(18)
    lead.paragraph_format.space_after = Pt(6)
    lead_run = lead.add_run(
        "This manual explains the main screens, the purpose of key buttons, and the usual click path for day-to-day document work."
    )
    apply_run_font(lead_run, size=11, color=MUTED_TEXT)

    doc.add_page_break()


def add_intro_tables(doc: Document) -> None:
    heading = doc.add_paragraph("How To Use This Manual", style="Heading 1")
    heading.paragraph_format.keep_with_next = True

    para = doc.add_paragraph()
    para.add_run(
        "The screenshots in this guide were captured from the live system on "
    )
    date_run = para.add_run(MANUAL_DATE)
    apply_run_font(date_run, size=11, bold=True)
    para.add_run(
        ". Your menu and available buttons may differ slightly depending on your assigned role, office, and permissions."
    )

    note = doc.add_table(rows=1, cols=1)
    note.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(note, width_inches=6.5, indent_dxa=0)
    set_table_borders(note, color=NOTE_BORDER, size="10")
    note_cell = note.cell(0, 0)
    shade_cell(note_cell, NOTE_FILL)
    set_cell_margins(note_cell, top=120, start=140, bottom=120, end=140)
    note_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    note_para = note_cell.paragraphs[0]
    note_para.paragraph_format.space_after = Pt(0)
    label_run = note_para.add_run("Important: ")
    apply_run_font(label_run, size=10.5, bold=True, color=HEADING_2_COLOR)
    body_run = note_para.add_run(
        "Use your assigned account only. Tracking IDs, routing history, and attachments should be handled according to office policy."
    )
    apply_run_font(body_run, size=10.5)

    doc.add_paragraph("Typical Daily Workflow", style="Heading 2")
    step_table = make_two_col_table(
        doc,
        headers=("Step", "What To Do"),
        rows=(
            ("1", "Sign in with your assigned username and password."),
            ("2", "Review the dashboard to see new workload, alerts, and due items."),
            ("3", "Open Create / Intake when you need to encode a new document."),
            ("4", "Submit the intake, then open the tracking slip to confirm the record."),
            ("5", "Use the archive/search page when you need to find completed records."),
            ("6", "Open Profile Settings to maintain your email and password."),
        ),
        widths=(STEP_COL_1, STEP_COL_2),
    )
    step_table.style = "Table Grid"


def make_two_col_table(
    doc: Document,
    headers: tuple[str, str],
    rows: Iterable[tuple[str, str]],
    widths: tuple[float, float],
):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, width_inches=6.5, indent_dxa=120)
    set_table_borders(table)
    set_column_widths(table, widths)

    hdr = table.rows[0]
    for idx, value in enumerate(headers):
        cell = hdr.cells[idx]
        shade_cell(cell, TABLE_HEADER_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(value)
        apply_run_font(run, size=10.5, bold=True, color=HEADING_2_COLOR)

    for left, right in rows:
        row = table.add_row()
        for idx, value in enumerate((left, right)):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(value)
            apply_run_font(run, size=10.5)

    return table


def add_image(doc: Document, image_path: Path, caption: str) -> None:
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(4)
    img_para.paragraph_format.space_after = Pt(4)
    img_para.add_run().add_picture(str(image_path), width=CONTENT_WIDTH)

    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.space_before = Pt(0)
    caption_para.paragraph_format.space_after = Pt(6)
    caption_run = caption_para.add_run(caption)
    apply_run_font(caption_run, size=9.5, color=MUTED_TEXT, italic=True)


def add_section(doc: Document, section_data: ManualSection) -> None:
    heading = doc.add_paragraph(section_data.title, style="Heading 1")
    heading.paragraph_format.keep_with_next = True

    for line in section_data.intro:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(line)
        apply_run_font(run, size=11)

    image_path = ASSET_DIR / section_data.image_name
    if not image_path.exists():
        raise FileNotFoundError(f"Missing screenshot asset: {image_path}")
    add_image(doc, image_path, section_data.figure_title)

    if section_data.steps:
        doc.add_paragraph("Walkthrough", style="Heading 2")
        make_two_col_table(
            doc,
            headers=("Step", "Action"),
            rows=section_data.steps,
            widths=(STEP_COL_1, STEP_COL_2),
        )

    doc.add_paragraph("Button Guide", style="Heading 2")
    make_two_col_table(
        doc,
        headers=("Callout", "Purpose"),
        rows=section_data.callouts,
        widths=(CALL_OUT_COL_1, CALL_OUT_COL_2),
    )


def add_closing(doc: Document) -> None:
    doc.add_paragraph("Support Tips", style="Heading 1")

    para = doc.add_paragraph()
    run = para.add_run(
        "When reporting a system issue, include the page name, the tracking ID if available, and a screenshot of the problem."
    )
    apply_run_font(run, size=11)

    table = make_two_col_table(
        doc,
        headers=("If You Need To...", "Best Action"),
        rows=(
            ("Reset access", "Use the password recovery page or contact your system administrator."),
            ("Trace a document", "Open the tracking slip and confirm the latest office and action."),
            ("Find a completed record", "Use the archive/search page with the date filter."),
            ("Change account details", "Open Profile Settings and verify changes with your current password."),
        ),
        widths=(Inches(2.1), Inches(4.4)),
    )
    table.style = "Table Grid"

    final_note = doc.add_paragraph()
    final_note.paragraph_format.space_before = Pt(10)
    final_note.paragraph_format.space_after = Pt(0)
    final_run = final_note.add_run(
        "End of manual. Use this guide together with your office procedures and the system's Support Center links."
    )
    apply_run_font(final_run, size=10.5, color=MUTED_TEXT, italic=True)


def build_manual() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_intro_tables(doc)
    for section in SECTIONS:
        add_section(doc, section)
    add_closing(doc)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_manual()
    print(path)
